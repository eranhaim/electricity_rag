"""
File processor for the Electricity RAG.

Extracts text from uploaded documents into LangChain ``Document`` objects
WITHOUT any LLM rewriting. Preserving the original text (Hebrew or otherwise)
is essential — an earlier "LLM optimization" step was translating Hebrew into
English and corrupting regulation numbers/values, which caused retrieval to
miss the actual content.

Each document is enriched with metadata:
    - source: original filename
    - page:   1-based page number (PDFs) or logical section (other formats)

A plain-text debug artifact is also written to ``data/processed/<name>.txt``
so a human can inspect exactly what the retriever will see.
"""

from __future__ import annotations

from pathlib import Path

from langchain.schema import Document

from backend.config import PROCESSED_DIR


def _load_pdf(file_path: Path) -> list[Document]:
    """Extract one Document per PDF page using PyPDFLoader (LangChain wrapper
    around pypdf). PyPDFLoader preserves per-page structure, unicode Hebrew,
    and does not modify the content."""
    from langchain_community.document_loaders import PyPDFLoader

    loader = PyPDFLoader(str(file_path))
    raw_docs = loader.load()

    docs: list[Document] = []
    for i, d in enumerate(raw_docs):
        text = (d.page_content or "").strip()
        if not text:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "page": i + 1,
                    "total_pages": len(raw_docs),
                },
            )
        )
    return docs


def _load_docx(file_path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    full_text = "\n".join(paragraphs)
    if not full_text.strip():
        return []
    return [
        Document(
            page_content=full_text,
            metadata={"source": file_path.name, "page": 1},
        )
    ]


def _load_txt(file_path: Path) -> list[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    if not text.strip():
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": file_path.name, "page": 1},
        )
    ]


def _load_xlsx(file_path: Path) -> list[Document]:
    import openpyxl

    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    docs: list[Document] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines = [f"Sheet: {sheet_name}"]
        for row in ws.iter_rows(values_only=True):
            lines.append("\t".join(str(c) if c is not None else "" for c in row))
        text = "\n".join(lines).strip()
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": file_path.name, "sheet": sheet_name},
                )
            )
    return docs


def _load_csv(file_path: Path) -> list[Document]:
    return _load_txt(file_path)


def load_file_documents(file_path: Path) -> list[Document]:
    """Dispatch to the right loader based on file extension.

    Returns a list of Documents preserving the ORIGINAL text and adding
    ``source`` + ``page`` metadata. NEVER translates or paraphrases content.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path)
    if suffix == ".docx":
        return _load_docx(file_path)
    if suffix in (".xlsx", ".xls"):
        return _load_xlsx(file_path)
    if suffix == ".csv":
        return _load_csv(file_path)
    return _load_txt(file_path)


def _write_debug_text(file_path: Path, docs: list[Document]) -> Path:
    """Write a human-readable plain-text artifact for transparency/debugging.
    The vector store is built from the Document objects directly — this file
    is NOT used for retrieval."""
    output_path = PROCESSED_DIR / (file_path.stem + ".txt")
    parts: list[str] = []
    for d in docs:
        page = d.metadata.get("page", d.metadata.get("sheet", "?"))
        parts.append(f"--- {file_path.name} | page/section {page} ---\n{d.page_content}")
    output_path.write_text("\n\n".join(parts), encoding="utf-8")
    return output_path


async def process_file(file_path: Path) -> Path:
    """Extract text from a file into Documents (no LLM rewriting), write a
    debug ``.txt`` artifact, and return the artifact path.

    The actual retrieval index is (re)built from the raw Documents by
    ``rag_pipeline.rebuild_vectorstore()``, which is called after this
    function by the upload handler.
    """
    docs = load_file_documents(file_path)
    if not docs:
        raise ValueError(f"No extractable text in {file_path.name}")
    return _write_debug_text(file_path, docs)
